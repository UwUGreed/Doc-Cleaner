"""
PII/PHI Sanitizer - Enterprise Privacy Tool
Detects, tokenizes, and rehydrates sensitive data with session-based self-destructing memory.
Supports: .txt, .pdf, .docx, .csv, .json, .xml, .html
"""

import customtkinter as ctk
from tkinterdnd2 import DND_FILES, TkinterDnD
import json
import os
import sys
import re
import secrets
import string
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple
import threading

from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer
from presidio_analyzer.nlp_engine import NlpEngineProvider

# PDF handling
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX handling
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class FileHandler:
    """Handles reading and writing multiple file formats."""
    
    @staticmethod
    def read_file(file_path: str) -> Tuple[str, str]:
        """
        Read file content and return (text, format).
        Returns: (content_text, file_format)
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return FileHandler._read_pdf(file_path), 'pdf'
        elif extension == '.docx':
            return FileHandler._read_docx(file_path), 'docx'
        elif extension in ['.txt', '.csv', '.json', '.xml', '.html', '.md']:
            return FileHandler._read_text(file_path), 'text'
        else:
            # Try as text
            return FileHandler._read_text(file_path), 'text'
    
    @staticmethod
    def _read_text(file_path: Path) -> str:
        """Read plain text file."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError(f"Could not decode file with common encodings")
    
    @staticmethod
    def _read_pdf(file_path: Path) -> str:
        """Read PDF file and extract text."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        text_content = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    text_content.append(f"--- Page {page_num + 1} (Error extracting text) ---\n")
        
        return "\n\n".join(text_content)
    
    @staticmethod
    def _read_docx(file_path: Path) -> str:
        """Read DOCX file and extract text."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        doc = Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_content.append(row_text)
        
        return "\n".join(text_content)
    
    @staticmethod
    def write_file(file_path: str, content: str, original_format: str):
        """
        Write content back to file in appropriate format.
        """
        file_path = Path(file_path)
        
        if original_format == 'pdf':
            FileHandler._write_pdf(file_path, content)
        elif original_format == 'docx':
            FileHandler._write_docx(file_path, content)
        else:
            FileHandler._write_text(file_path, content)
    
    @staticmethod
    def _write_text(file_path: Path, content: str):
        """Write plain text file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    @staticmethod
    def _write_pdf(file_path: Path, content: str):
        """Write PDF file (using reportlab)."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split content into paragraphs
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    # Handle page markers
                    if para.startswith('--- Page'):
                        story.append(Spacer(1, 0.2*inch))
                        story.append(Paragraph(f"<b>{para}</b>", styles['Heading2']))
                        story.append(Spacer(1, 0.1*inch))
                    else:
                        # Escape special characters for reportlab
                        safe_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(safe_para, styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            
        except ImportError:
            # Fallback: save as text with .pdf extension
            FileHandler._write_text(file_path, content)
    
    @staticmethod
    def _write_docx(file_path: Path, content: str):
        """Write DOCX file."""
        if not DOCX_AVAILABLE:
            # Fallback: save as text
            FileHandler._write_text(file_path, content)
            return
        
        doc = Document()
        
        # Split content into paragraphs
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                # Detect table rows (contains |)
                if ' | ' in para:
                    # This was likely from a table - add as paragraph for simplicity
                    doc.add_paragraph(para)
                else:
                    doc.add_paragraph(para)
        
        doc.save(str(file_path))


class SessionVault:
    """Manages session-based token mappings with self-destruct capability."""
    
    def __init__(self):
        self.vault_path = Path(os.getenv('APPDATA')) / 'PIISanitizer' / 'session_vault.json'
        self.vault_path.parent.mkdir(parents=True, exist_ok=True)
        self.mappings: Dict[str, Dict[str, str]] = self._load_vault()
        
    def _load_vault(self) -> Dict[str, Dict[str, str]]:
        """Load existing vault or create new one."""
        if self.vault_path.exists():
            try:
                with open(self.vault_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {}
        return {}
    
    def _save_vault(self):
        """Persist vault to disk."""
        with open(self.vault_path, 'w', encoding='utf-8') as f:
            json.dump(self.mappings, f, indent=2, ensure_ascii=False)
    
    def generate_token(self, entity_type: str) -> str:
        """Generate a cryptographically secure unique token."""
        # 4-character alphanumeric suffix for uniqueness
        suffix = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(4))
        return f"<{entity_type}_{suffix}>"
    
    def store_mapping(self, session_id: str, token: str, original_value: str):
        """Store token-to-value mapping for a session."""
        if session_id not in self.mappings:
            self.mappings[session_id] = {}
        self.mappings[session_id][token] = original_value
        self._save_vault()
    
    def get_mapping(self, session_id: str) -> Dict[str, str]:
        """Retrieve all mappings for a session."""
        return self.mappings.get(session_id, {})
    
    def delete_session(self, session_id: str):
        """Delete a specific session's mappings (The Shredder Protocol)."""
        if session_id in self.mappings:
            del self.mappings[session_id]
            self._save_vault()
    
    def shred_vault(self):
        """Complete vault destruction on app close."""
        if self.vault_path.exists():
            self.vault_path.unlink()


class PresidioEngine:
    """Enhanced Presidio engine with custom recognizers."""
    
    def __init__(self, log_callback):
        self.log_callback = log_callback
        self.analyzer = None
        self._initialize_engine()
    
    def _initialize_engine(self):
        """Initialize Presidio with custom recognizers."""
        self.log_callback("🔧 Initializing Presidio Analyzer Engine...")
        
        # Configure NLP engine with SpaCy
        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}],
        }
        
        try:
            provider = NlpEngineProvider(nlp_configuration=configuration)
            nlp_engine = provider.create_engine()
            self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
            self.log_callback("✅ SpaCy model loaded successfully")
        except:
            # Fallback to small model
            self.log_callback("⚠️ en_core_web_lg not found, using en_core_web_sm")
            configuration["models"][0]["model_name"] = "en_core_web_sm"
            provider = NlpEngineProvider(nlp_configuration=configuration)
            nlp_engine = provider.create_engine()
            self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
        
        # Add custom recognizers
        self._add_insurance_policy_recognizer()
        self._add_medical_record_recognizer()
        self._add_drivers_license_recognizer()
        self._add_health_plan_recognizer()
        
        self.log_callback("✅ Presidio Engine Ready with Custom Recognizers")
    
    def _add_insurance_policy_recognizer(self):
        """Add recognizer for insurance policy numbers (POL-12345678)."""
        policy_pattern = Pattern(
            name="insurance_policy_pattern",
            regex=r"\b(?:POL|POLICY|INS|P)-?\d{6,12}\b",
            score=0.85
        )
        policy_recognizer = PatternRecognizer(
            supported_entity="INSURANCE_POLICY",
            patterns=[policy_pattern]
        )
        self.analyzer.registry.add_recognizer(policy_recognizer)
        self.log_callback("  ➕ Insurance Policy Number recognizer added")
    
    def _add_medical_record_recognizer(self):
        """Add recognizer for Medical Record Numbers (MRN)."""
        mrn_pattern = Pattern(
            name="mrn_pattern",
            regex=r"\b(?:MRN|MR|MEDICAL.?RECORD|PATIENT.?ID)[:\s#-]*(\d{6,10})\b",
            score=0.9
        )
        mrn_recognizer = PatternRecognizer(
            supported_entity="MEDICAL_RECORD_NUMBER",
            patterns=[mrn_pattern]
        )
        self.analyzer.registry.add_recognizer(mrn_recognizer)
        self.log_callback("  ➕ Medical Record Number (MRN) recognizer added")
    
    def _add_drivers_license_recognizer(self):
        """Add recognizer for US Driver's License numbers."""
        dl_pattern = Pattern(
            name="drivers_license_pattern",
            regex=r"\b(?:DL|DRIVER.?LICENSE)[:\s#-]*([A-Z0-9]{6,12})\b",
            score=0.85
        )
        dl_recognizer = PatternRecognizer(
            supported_entity="DRIVERS_LICENSE",
            patterns=[dl_pattern]
        )
        self.analyzer.registry.add_recognizer(dl_recognizer)
        self.log_callback("  ➕ Driver's License recognizer added")
    
    def _add_health_plan_recognizer(self):
        """Add recognizer for Health Plan ID numbers."""
        health_plan_pattern = Pattern(
            name="health_plan_pattern",
            regex=r"\b(?:HEALTH.?PLAN|INSURANCE.?ID|MEMBER.?ID)[:\s#-]*([A-Z0-9]{8,15})\b",
            score=0.85
        )
        health_plan_recognizer = PatternRecognizer(
            supported_entity="HEALTH_PLAN_ID",
            patterns=[health_plan_pattern]
        )
        self.analyzer.registry.add_recognizer(health_plan_recognizer)
        self.log_callback("  ➕ Health Plan ID recognizer added")
    
    def analyze_text(self, text: str) -> List:
        """Analyze text and return detected PII/PHI entities."""
        # Full entity list for comprehensive detection
        entities = [
            "PERSON",  # Full names
            "PHONE_NUMBER",
            "EMAIL_ADDRESS",
            "LOCATION",  # Addresses
            "US_SSN",
            "CREDIT_CARD",
            "US_BANK_NUMBER",
            "IBAN_CODE",
            "DATE_TIME",  # Birthdates
            "MEDICAL_LICENSE",
            "US_PASSPORT",
            "INSURANCE_POLICY",  # Custom
            "MEDICAL_RECORD_NUMBER",  # Custom
            "DRIVERS_LICENSE",  # Custom
            "HEALTH_PLAN_ID",  # Custom
        ]
        
        results = self.analyzer.analyze(
            text=text,
            entities=entities,
            language='en',
            allow_list=None
        )
        
        # Sort by start position (reverse order for replacement)
        results.sort(key=lambda x: x.start, reverse=True)
        return results


class PIISanitizerApp(TkinterDnD.Tk):
    """Main application window with drag-and-drop support."""
    
    def __init__(self):
        super().__init__()
        
        # Initialize session vault
        self.vault = SessionVault()
        
        # App configuration
        self.title("PII/PHI Sanitizer - Enterprise Privacy Tool")
        self.geometry("1000x700")
        self.configure(bg="#1a1a1a")
        
        # Set theme
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # Initialize Presidio engine in background
        self.presidio_engine = None
        self.engine_ready = False
        
        self.setup_ui()
        self.initialize_engine()
        
        # Bind app close to vault shredding
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def setup_ui(self):
        """Create the user interface."""
        
        # Header Frame
        header_frame = ctk.CTkFrame(self, fg_color="#2b2b2b", corner_radius=0)
        header_frame.pack(fill="x", padx=0, pady=0)
        
        header_label = ctk.CTkLabel(
            header_frame,
            text="🔒 PII/PHI Sanitizer",
            font=ctk.CTkFont(size=24, weight="bold"),
            text_color="#ffffff"
        )
        header_label.pack(side="left", padx=20, pady=15)
        
        # Theme toggle
        self.theme_switch = ctk.CTkSwitch(
            header_frame,
            text="Light Mode",
            command=self.toggle_theme,
            font=ctk.CTkFont(size=12)
        )
        self.theme_switch.pack(side="right", padx=20, pady=15)
        
        # Main content area
        content_frame = ctk.CTkFrame(self, fg_color="transparent")
        content_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Drop Zone
        self.drop_frame = ctk.CTkFrame(
            content_frame,
            fg_color="#2b2b2b",
            border_width=3,
            border_color="#3b8ed0",
            corner_radius=10
        )
        self.drop_frame.pack(fill="x", pady=(0, 20))
        
        drop_label = ctk.CTkLabel(
            self.drop_frame,
            text="📁 Drag & Drop Files Here\n\nSupports: .txt, .pdf, .docx, .csv, .json, .xml, .html",
            font=ctk.CTkFont(size=18),
            text_color="#aaaaaa"
        )
        drop_label.pack(pady=40)
        
        # Enable drag and drop
        self.drop_frame.drop_target_register(DND_FILES)
        self.drop_frame.dnd_bind('<<Drop>>', self.on_file_drop)
        
        # Action Buttons Frame
        button_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
        button_frame.pack(fill="x", pady=(0, 20))
        
        self.sanitize_btn = ctk.CTkButton(
            button_frame,
            text="🔍 Sanitize File (Remove PII/PHI)",
            font=ctk.CTkFont(size=14, weight="bold"),
            height=40,
            command=self.select_file_sanitize,
            state="disabled"
        )
        self.sanitize_btn.pack(side="left", padx=(0, 10), expand=True, fill="x")
        
        self.rehydrate_btn = ctk.CTkButton(
            button_frame,
            text="💧 Rehydrate File (Restore Original Data)",
            font=ctk.CTkFont(size=14, weight="bold"),
            height=40,
            command=self.select_file_rehydrate,
            state="disabled",
            fg_color="#2c8c2c",
            hover_color="#236e23"
        )
        self.rehydrate_btn.pack(side="left", expand=True, fill="x")
        
        # Log Output
        log_label = ctk.CTkLabel(
            content_frame,
            text="📋 Activity Log",
            font=ctk.CTkFont(size=14, weight="bold"),
            anchor="w"
        )
        log_label.pack(fill="x", pady=(0, 5))
        
        self.log_text = ctk.CTkTextbox(
            content_frame,
            font=ctk.CTkFont(family="Consolas", size=11),
            fg_color="#1e1e1e",
            text_color="#00ff00",
            wrap="word"
        )
        self.log_text.pack(fill="both", expand=True)
        
        # Status Bar
        self.status_label = ctk.CTkLabel(
            self,
            text="⏳ Initializing Presidio Engine...",
            font=ctk.CTkFont(size=11),
            anchor="w",
            fg_color="#2b2b2b",
            corner_radius=0
        )
        self.status_label.pack(fill="x", padx=0, pady=0)
        
        self.log("🚀 Application started. Waiting for Presidio initialization...")
    
    def initialize_engine(self):
        """Initialize Presidio engine in background thread."""
        def init_thread():
            self.presidio_engine = PresidioEngine(self.log)
            self.engine_ready = True
            self.after(0, self._enable_buttons)
        
        thread = threading.Thread(target=init_thread, daemon=True)
        thread.start()
    
    def _enable_buttons(self):
        """Enable buttons after engine initialization."""
        self.sanitize_btn.configure(state="normal")
        self.rehydrate_btn.configure(state="normal")
        self.status_label.configure(text="✅ Ready to sanitize files")
        self.log("🟢 System Ready. You can now sanitize or rehydrate files.")
    
    def toggle_theme(self):
        """Toggle between light and dark mode."""
        if self.theme_switch.get():
            ctk.set_appearance_mode("light")
            self.log("☀️ Switched to Light Mode")
        else:
            ctk.set_appearance_mode("dark")
            self.log("🌙 Switched to Dark Mode")
    
    def log(self, message: str):
        """Add message to log."""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert("end", f"[{timestamp}] {message}\n")
        self.log_text.see("end")
        self.update_idletasks()
    
    def on_file_drop(self, event):
        """Handle file drop event."""
        if not self.engine_ready:
            self.log("⚠️ Please wait for engine initialization to complete.")
            return
        
        file_path = event.data.strip('{}')
        self.log(f"📥 File dropped: {Path(file_path).name}")
        self.sanitize_file(file_path)
    
    def select_file_sanitize(self):
        """Open file dialog to select file for sanitization."""
        from tkinter import filedialog
        file_path = filedialog.askopenfilename(
            title="Select file to sanitize",
            filetypes=[
                ("All supported files", "*.txt *.pdf *.docx *.csv *.json *.xml *.html"),
                ("Text files", "*.txt"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("CSV files", "*.csv"),
                ("JSON files", "*.json"),
                ("XML files", "*.xml"),
                ("HTML files", "*.html"),
                ("All files", "*.*")
            ]
        )
        if file_path:
            self.sanitize_file(file_path)
    
    def select_file_rehydrate(self):
        """Open file dialog to select file for rehydration."""
        from tkinter import filedialog
        file_path = filedialog.askopenfilename(
            title="Select sanitized file to rehydrate",
            filetypes=[
                ("All supported files", "*.txt *.pdf *.docx *.csv *.json *.xml *.html"),
                ("Text files", "*.txt"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("CSV files", "*.csv"),
                ("JSON files", "*.json"),
                ("XML files", "*.xml"),
                ("HTML files", "*.html"),
                ("All files", "*.*")
            ]
        )
        if file_path:
            self.rehydrate_file(file_path)
    
    def sanitize_file(self, file_path: str):
        """Sanitize a file by replacing PII/PHI with tokens."""
        try:
            file_name = Path(file_path).name
            self.log(f"🔍 Scanning: {file_name}")
            self.status_label.configure(text="⏳ Analyzing file for PII/PHI...")
            
            # Read file using FileHandler
            try:
                original_text, file_format = FileHandler.read_file(file_path)
                self.log(f"📄 Detected format: {file_format.upper()}")
            except ImportError as e:
                self.log(f"❌ Missing library: {str(e)}")
                self.status_label.configure(text="❌ Missing required library")
                return
            except Exception as e:
                self.log(f"❌ Failed to read file: {str(e)}")
                self.status_label.configure(text="❌ File read error")
                return
            
            # Analyze with Presidio
            results = self.presidio_engine.analyze_text(original_text)
            
            if not results:
                self.log("✅ No PII/PHI detected. File is clean.")
                self.status_label.configure(text="✅ No sensitive data found")
                return
            
            # Generate session ID
            session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{secrets.token_hex(4)}"
            
            # Count entities by type
            entity_counts = {}
            sanitized_text = original_text
            
            # Replace detected entities with tokens
            for result in results:
                entity_type = result.entity_type
                original_value = original_text[result.start:result.end]
                
                # Generate unique token
                token = self.vault.generate_token(entity_type)
                
                # Store mapping
                self.vault.store_mapping(session_id, token, original_value)
                
                # Replace in text
                sanitized_text = (
                    sanitized_text[:result.start] + 
                    token + 
                    sanitized_text[result.end:]
                )
                
                # Count entities
                entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
            
            # Save sanitized file
            original_path = Path(file_path)
            output_path = original_path.parent / f"{original_path.stem}_SANITIZED{original_path.suffix}"
            
            # Write using FileHandler
            try:
                FileHandler.write_file(str(output_path), sanitized_text, file_format)
            except Exception as e:
                self.log(f"⚠️ Could not write as {file_format.upper()}, saving as text")
                output_path = original_path.parent / f"{original_path.stem}_SANITIZED.txt"
                FileHandler.write_file(str(output_path), sanitized_text, 'text')
            
            # Save session ID and format to metadata file
            metadata_path = original_path.parent / f"{original_path.stem}_SANITIZED.session"
            metadata = {
                'session_id': session_id,
                'original_format': file_format,
                'original_extension': original_path.suffix
            }
            with open(metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            # Log results
            self.log("✅ SANITIZATION COMPLETE")
            self.log(f"📊 Detected and Scrubbed:")
            for entity_type, count in sorted(entity_counts.items()):
                self.log(f"   • {count}x {entity_type}")
            self.log(f"💾 Safe file saved: {output_path.name}")
            self.log(f"🔑 Session ID: {session_id}")
            self.status_label.configure(text=f"✅ Scrubbed {sum(entity_counts.values())} sensitive items")
            
        except Exception as e:
            self.log(f"❌ Error: {str(e)}")
            self.status_label.configure(text="❌ Sanitization failed")
            import traceback
            self.log(f"Debug: {traceback.format_exc()}")
    
    def rehydrate_file(self, file_path: str):
        """Rehydrate a sanitized file by restoring original PII/PHI."""
        try:
            file_name = Path(file_path).name
            self.log(f"💧 Rehydrating: {file_name}")
            self.status_label.configure(text="⏳ Restoring original data...")
            
            # Look for session file
            session_path = Path(file_path).parent / f"{Path(file_path).stem}.session"
            if not session_path.exists():
                self.log("❌ Session file not found. Cannot rehydrate.")
                self.log("   Make sure the .session file is in the same directory.")
                self.status_label.configure(text="❌ Session file missing")
                return
            
            # Read session metadata
            try:
                with open(session_path, 'r') as f:
                    content = f.read().strip()
                    # Check if it's JSON (new format) or plain text (old format)
                    if content.startswith('{'):
                        metadata = json.loads(content)
                        session_id = metadata['session_id']
                        file_format = metadata.get('original_format', 'text')
                        original_extension = metadata.get('original_extension', Path(file_path).suffix)
                    else:
                        # Old format - just session ID
                        session_id = content
                        file_format = 'text'
                        original_extension = Path(file_path).suffix
            except Exception as e:
                self.log(f"❌ Failed to read session file: {str(e)}")
                self.status_label.configure(text="❌ Session file error")
                return
            
            # Get mappings
            mappings = self.vault.get_mapping(session_id)
            if not mappings:
                self.log("❌ No mappings found for this session.")
                self.status_label.configure(text="❌ Session data not found")
                return
            
            # Read sanitized file
            try:
                sanitized_text, detected_format = FileHandler.read_file(file_path)
            except Exception as e:
                self.log(f"❌ Failed to read sanitized file: {str(e)}")
                self.status_label.configure(text="❌ File read error")
                return
            
            # Replace tokens with original values
            rehydrated_text = sanitized_text
            restore_count = 0
            for token, original_value in mappings.items():
                if token in rehydrated_text:
                    rehydrated_text = rehydrated_text.replace(token, original_value)
                    restore_count += 1
            
            # Save rehydrated file
            original_path = Path(file_path)
            output_path = original_path.parent / f"{original_path.stem}_RESTORED{original_extension}"
            
            # Write using FileHandler
            try:
                FileHandler.write_file(str(output_path), rehydrated_text, file_format)
            except Exception as e:
                self.log(f"⚠️ Could not write as {file_format.upper()}, saving as text")
                output_path = original_path.parent / f"{original_path.stem}_RESTORED.txt"
                FileHandler.write_file(str(output_path), rehydrated_text, 'text')
            
            # THE SHREDDER PROTOCOL: Delete session data
            self.vault.delete_session(session_id)
            session_path.unlink()  # Delete session file
            
            self.log("✅ REHYDRATION COMPLETE")
            self.log(f"🔓 Restored {restore_count} sensitive items")
            self.log(f"💾 Restored file saved: {output_path.name}")
            self.log("🔥 THE SHREDDER PROTOCOL ACTIVATED")
            self.log(f"   • Session '{session_id}' deleted from vault")
            self.log(f"   • Session file removed from disk")
            self.status_label.configure(text=f"✅ Restored {restore_count} items & shredded session")
            
        except Exception as e:
            self.log(f"❌ Error: {str(e)}")
            self.status_label.configure(text="❌ Rehydration failed")
            import traceback
            self.log(f"Debug: {traceback.format_exc()}")
    
    def on_closing(self):
        """Handle application closing - shred the vault."""
        self.log("🔥 Application closing - initiating vault destruction...")
        self.vault.shred_vault()
        self.log("✅ Vault completely wiped. All session data destroyed.")
        self.destroy()


def main():
    """Main entry point."""
    app = PIISanitizerApp()
    app.mainloop()


if __name__ == "__main__":
    main()
