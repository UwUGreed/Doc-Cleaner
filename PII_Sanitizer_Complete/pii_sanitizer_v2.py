"""
PII/PHI Sanitizer - Enterprise Privacy Tool
Detects, tokenizes, and rehydrates sensitive data with session-based self-destructing memory.
Supports: .txt, .pdf, .docx, .csv, .json, .xml, .html
"""

import customtkinter as ctk
import tkinter as tk
from tkinterdnd2 import DND_FILES, TkinterDnD
import json
import os
import sys
import re
import secrets
import string
import difflib
import unicodedata
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Set, Tuple
import threading

from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer
from presidio_analyzer.nlp_engine import NlpEngineProvider

# PDF handling — prefer PyMuPDF (fitz) for proper in-place redaction
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX handling
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class FileHandler:
    """Handles reading and writing multiple file formats."""

    @staticmethod
    def read_file(file_path: str) -> Tuple[str, str]:
        """Read file content and return (text, format)."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return FileHandler._read_pdf(file_path), 'pdf'
        elif extension == '.docx':
            return FileHandler._read_docx(file_path), 'docx'
        elif extension in ['.txt', '.csv', '.json', '.xml', '.html', '.md']:
            return FileHandler._read_text(file_path), 'text'
        else:
            return FileHandler._read_text(file_path), 'text'

    @staticmethod
    def _read_text(file_path: Path) -> str:
        """Read plain text file."""
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError("Could not decode file with common encodings")

    @staticmethod
    def _read_pdf(file_path: Path) -> str:
        """Read PDF. Prefers PyMuPDF for accuracy; falls back to PyPDF2."""
        if PYMUPDF_AVAILABLE:
            return FileHandler._read_pdf_fitz(file_path)
        elif PDF_AVAILABLE:
            return FileHandler._read_pdf_pypdf2(file_path)
        else:
            raise ImportError("No PDF library found. Install with: pip install pymupdf")

    @staticmethod
    def _read_pdf_fitz(file_path: Path) -> str:
        """Extract PDF text using PyMuPDF (consistent with fitz search_for)."""
        doc = fitz.open(str(file_path))
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append(f"--- Page {i + 1} ---\n{text}")
        doc.close()
        return "\n\n".join(pages)

    @staticmethod
    def _read_pdf_pypdf2(file_path: Path) -> str:
        """Fallback PDF reader using PyPDF2."""
        text_content = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(f"--- Page {i + 1} ---\n{text}")
                except Exception:
                    text_content.append(f"--- Page {i + 1} (Error extracting text) ---\n")
        return "\n\n".join(text_content)

    @staticmethod
    def _read_docx(file_path: Path) -> str:
        """Read DOCX file and extract text."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        doc = Document(file_path)
        parts = []
        for para in FileHandler._iter_docx_paragraphs(doc):
            if para.text.strip():
                parts.append(para.text)
        return "\n".join(parts)

    # ------------------------------------------------------------------ #
    #  PDF write paths (require PyMuPDF)                                  #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _rects_overlap(r1: "fitz.Rect", r2: "fitz.Rect") -> bool:
        """Return True if two rects share any area."""
        return not (r1.x1 <= r2.x0 or r2.x1 <= r1.x0 or
                    r1.y1 <= r2.y0 or r2.y1 <= r1.y0)

    @staticmethod
    def _tighten_redaction_rect(rect: "fitz.Rect") -> "fitz.Rect":
        """
        PyMuPDF search rectangles can include descender padding that bleeds into
        the next line. Redacting the padded box can erase unrelated nearby text.
        """
        tightened = fitz.Rect(rect)
        if tightened.height > 4:
            tightened.y0 += tightened.height * 0.04
            tightened.y1 -= tightened.height * 0.16
        return tightened

    @staticmethod
    def _same_text_line(r1: "fitz.Rect", r2: "fitz.Rect") -> bool:
        mid1 = (r1.y0 + r1.y1) / 2
        mid2 = (r2.y0 + r2.y1) / 2
        return abs(mid1 - mid2) <= max(r1.height, r2.height) * 0.60

    @staticmethod
    def _expanded_line_rect(page: "fitz.Page",
                            rect: "fitz.Rect",
                            page_rects: List["fitz.Rect"]) -> "fitz.Rect":
        """Give long restored values room without dropping into the next line."""
        right_margin = page.rect.x1 - 24
        same_line_next = [
            other.x0 for other in page_rects
            if other is not rect and other.x0 > rect.x0 and
            FileHandler._same_text_line(rect, other)
        ]
        if same_line_next:
            return rect
        x1 = max(right_margin, rect.x1)

        expanded = fitz.Rect(rect.x0, rect.y0, x1, rect.y1)
        expanded.y0 -= rect.height * 0.08
        expanded.y1 += rect.height * 0.04
        return expanded

    @staticmethod
    def _fit_fontsize(text: str,
                      rect: "fitz.Rect",
                      fontname: str,
                      preferred_size: float,
                      min_size: float = 5.5) -> float:
        if not text:
            return preferred_size
        available_width = max(rect.width - 1, 1)
        fontsize = preferred_size
        while fontsize > min_size:
            try:
                width = fitz.get_text_length(text, fontname=fontname, fontsize=fontsize)
            except Exception:
                width = len(text) * fontsize * 0.5
            if width <= available_width and fontsize <= max(rect.height * 0.92, min_size):
                return fontsize
            fontsize -= 0.25
        return min_size

    @staticmethod
    def _insert_fitted_text(page: "fitz.Page",
                            rect: "fitz.Rect",
                            text: str,
                            preferred_fontsize: float,
                            color: tuple):
        fontname = "helv"
        fontsize = FileHandler._fit_fontsize(text, rect, fontname, preferred_fontsize)
        baseline = fitz.Point(rect.x0, rect.y0 + (rect.height * 0.72))
        page.insert_text(
            baseline, text,
            fontsize=fontsize,
            fontname=fontname,
            color=color,
        )

    @staticmethod
    def _pdf_replace(doc: "fitz.Document",
                     search_map: dict,
                     insert_fontsize: float,
                     insert_color: tuple,
                     expand_for_rehydration: bool = False) -> int:
        """
        Core PDF text-replacement loop shared by sanitize and rehydrate.

        For every key in search_map:
          1. search_for(key)  — locate all occurrences on each page
          2. add_redact_annot(fill=white) — wipe the original bytes from the stream
          3. apply_redactions()
          4. insert_text(baseline_point, value) — write the replacement into the stream

        insert_text (point-based) is used instead of insert_textbox because
        insert_textbox silently drops text that overflows the rect, leaving
        nothing in the text layer — which is exactly why rehydration was
        returning 0 hits.
        """
        hit_count = 0
        for page in doc:
            replacements: List[Tuple] = []  # (fitz.Rect, replacement_text)

            for search_key, replacement in sorted(
                    search_map.items(), key=lambda item: len(item[0]), reverse=True):
                if not search_key or not search_key.strip():
                    continue
                for rect in page.search_for(search_key):
                    r = FileHandler._tighten_redaction_rect(fitz.Rect(rect))
                    # Skip if this rect overlaps one we've already queued
                    # (prevents double-redacting the same area)
                    if any(FileHandler._rects_overlap(r, queued)
                           for queued, _ in replacements):
                        continue
                    replacements.append((r, replacement))
                    hit_count += 1

            if not replacements:
                continue

            original_rects = [rect for rect, _ in replacements]
            prepared_replacements = []
            for rect, text in replacements:
                write_rect = (
                    FileHandler._expanded_line_rect(page, rect, original_rects)
                    if expand_for_rehydration else rect
                )
                prepared_replacements.append((write_rect, text))
                page.add_redact_annot(write_rect, fill=(1, 1, 1))

            try:
                page.apply_redactions(images=0)   # 0 = PDF_REDACT_IMAGE_NONE
            except TypeError:
                page.apply_redactions()

            for rect, text in prepared_replacements:
                FileHandler._insert_fitted_text(
                    page, rect, text, insert_fontsize, insert_color)

        return hit_count

    @staticmethod
    def write_pdf_sanitized(original_path: str, output_path: str,
                             value_to_token: dict) -> int:
        """Replace PII values with tokens; tokens written in small red text."""
        doc = fitz.open(str(original_path))
        count = FileHandler._pdf_replace(
            doc,
            search_map=value_to_token,
            insert_fontsize=7,
            insert_color=(0.75, 0, 0),
        )
        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()
        return count

    @staticmethod
    def write_pdf_masked(original_path: str, output_path: str,
                         value_to_mask: dict) -> int:
        """Replace PII values with irreversible masks."""
        doc = fitz.open(str(original_path))
        count = FileHandler._pdf_replace(
            doc,
            search_map=value_to_mask,
            insert_fontsize=8,
            insert_color=(0, 0, 0),
        )
        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()
        return count

    @staticmethod
    def write_pdf_rehydrated(sanitized_path: str, output_path: str,
                              token_to_value: dict) -> int:
        """Replace tokens with original values; written in standard black text."""
        doc = fitz.open(str(sanitized_path))
        count = FileHandler._pdf_replace(
            doc,
            search_map=token_to_value,
            insert_fontsize=9,
            insert_color=(0, 0, 0),
            expand_for_rehydration=True,
        )
        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()
        return count

    @staticmethod
    def _iter_docx_paragraphs(container):
        for paragraph in getattr(container, "paragraphs", []):
            yield paragraph
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    yield from FileHandler._iter_docx_paragraphs(cell)
        for section in getattr(container, "sections", []):
            yield from FileHandler._iter_docx_paragraphs(section.header)
            yield from FileHandler._iter_docx_paragraphs(section.footer)
            yield from FileHandler._iter_docx_paragraphs(section.first_page_header)
            yield from FileHandler._iter_docx_paragraphs(section.first_page_footer)
            yield from FileHandler._iter_docx_paragraphs(section.even_page_header)
            yield from FileHandler._iter_docx_paragraphs(section.even_page_footer)

    @staticmethod
    def _replace_in_docx_paragraph(paragraph, replacements: dict) -> int:
        runs = paragraph.runs
        if not runs:
            return 0

        full_text = "".join(run.text for run in runs)
        if not full_text:
            return 0

        matches = []
        occupied = [False] * len(full_text)
        for search_text, replacement in sorted(
                replacements.items(), key=lambda item: len(item[0]), reverse=True):
            if not search_text:
                continue
            start = full_text.find(search_text)
            while start != -1:
                end = start + len(search_text)
                if not any(occupied[start:end]):
                    matches.append((start, end, replacement))
                    for i in range(start, end):
                        occupied[i] = True
                start = full_text.find(search_text, start + 1)

        if not matches:
            return 0

        for start, end, replacement in sorted(matches, key=lambda item: item[0], reverse=True):
            positions = []
            for run_index, run in enumerate(runs):
                positions.extend((run_index, offset) for offset, _ in enumerate(run.text))

            if start >= len(positions) or end - 1 >= len(positions):
                continue

            start_run_index, start_offset = positions[start]
            end_run_index, end_offset = positions[end - 1]
            end_offset += 1

            if start_run_index == end_run_index:
                run = runs[start_run_index]
                run.text = run.text[:start_offset] + replacement + run.text[end_offset:]
                continue

            start_run = runs[start_run_index]
            end_run = runs[end_run_index]
            start_run.text = (
                start_run.text[:start_offset] +
                replacement +
                end_run.text[end_offset:]
            )
            for run_index in range(start_run_index + 1, end_run_index + 1):
                runs[run_index].text = ""

        return len(matches)

    @staticmethod
    def _replace_docx(original_path: str, output_path: str, replacements: dict) -> int:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        doc = Document(original_path)
        count = 0
        for paragraph in FileHandler._iter_docx_paragraphs(doc):
            count += FileHandler._replace_in_docx_paragraph(paragraph, replacements)
        doc.save(str(output_path))
        return count

    @staticmethod
    def write_docx_sanitized(original_path: str, output_path: str,
                             value_to_token: dict) -> int:
        return FileHandler._replace_docx(original_path, output_path, value_to_token)

    @staticmethod
    def write_docx_masked(original_path: str, output_path: str,
                          value_to_mask: dict) -> int:
        return FileHandler._replace_docx(original_path, output_path, value_to_mask)

    @staticmethod
    def write_docx_rehydrated(sanitized_path: str, output_path: str,
                              token_to_value: dict) -> int:
        return FileHandler._replace_docx(sanitized_path, output_path, token_to_value)

    @staticmethod
    def write_file(file_path: str, content: str, original_format: str):
        """Write content back to file in appropriate format."""
        file_path = Path(file_path)
        if original_format == 'docx':
            FileHandler._write_docx(file_path, content)
        else:
            FileHandler._write_text(file_path, content)

    @staticmethod
    def _write_text(file_path: Path, content: str):
        """Write plain text file."""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

    @staticmethod
    def _write_docx(file_path: Path, content: str):
        """Write DOCX file."""
        if not DOCX_AVAILABLE:
            FileHandler._write_text(file_path, content)
            return
        doc = Document()
        for para in content.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
        doc.save(str(file_path))


class SessionVault:
    """Manages session-based token mappings with self-destruct capability."""

    def __init__(self):
        self.vault_path = Path(os.getenv('APPDATA')) / 'PIISanitizer' / 'session_vault.json'
        self.vault_path.parent.mkdir(parents=True, exist_ok=True)
        self.mappings: Dict[str, Dict[str, str]] = self._load_vault()
        self.token_counter = self._next_token_counter()

    def _next_token_counter(self) -> int:
        highest = 0
        for mapping in self.mappings.values():
            for token in mapping:
                match = re.match(r"\bToken_(\d+)\b", token, flags=re.IGNORECASE)
                if match:
                    highest = max(highest, int(match.group(1)))
                    continue
                match = re.match(r"\bTOKEN_(\d{6})_[A-Z][A-Z0-9_]*\b", token)
                if match:
                    highest = max(highest, int(match.group(1)))
        return highest + 1

    def _load_vault(self) -> Dict[str, Dict[str, str]]:
        if self.vault_path.exists():
            try:
                with open(self.vault_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _save_vault(self):
        with open(self.vault_path, 'w', encoding='utf-8') as f:
            json.dump(self.mappings, f, indent=2, ensure_ascii=False)

    def generate_token(self, entity_type: str) -> str:
        token = f"Token_{self.token_counter}"
        self.token_counter += 1
        return token

    def store_mapping(self, session_id: str, token: str, original_value: str):
        if session_id not in self.mappings:
            self.mappings[session_id] = {}
        self.mappings[session_id][token] = original_value
        self._save_vault()

    def get_mapping(self, session_id: str) -> Dict[str, str]:
        return self.mappings.get(session_id, {})

    def find_session_by_tokens(self, tokens: Set[str]) -> Tuple[Optional[str], int, bool]:
        """Find the vault session whose stored tokens best match a file."""
        if not tokens:
            return None, 0, False

        matches = []
        for session_id, mapping in self.mappings.items():
            match_count = len(tokens.intersection(mapping.keys()))
            if match_count:
                matches.append((session_id, match_count))

        if not matches:
            return None, 0, False

        matches.sort(key=lambda item: item[1], reverse=True)
        best_session_id, best_count = matches[0]
        ambiguous = len(matches) > 1 and matches[1][1] == best_count
        return best_session_id, best_count, ambiguous

    @staticmethod
    def token_match_key(token: str) -> str:
        normalized = unicodedata.normalize("NFKD", token or "")
        ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
        return re.sub(r"[^A-Z0-9]", "", ascii_text.upper())

    def find_session_by_token_candidates(
            self, candidates: Set[str]) -> Tuple[Optional[str], int, bool]:
        """Find a session even when tokens were cosmetically changed."""
        if not candidates:
            return None, 0, False

        candidate_keys = {self.token_match_key(candidate) for candidate in candidates}
        matches = []
        for session_id, mapping in self.mappings.items():
            token_keys = {self.token_match_key(token) for token in mapping.keys()}
            exact_count = len(candidate_keys.intersection(token_keys))
            fuzzy_count = 0
            for candidate_key in candidate_keys - token_keys:
                if not candidate_key:
                    continue
                best_ratio = max(
                    (difflib.SequenceMatcher(None, candidate_key, token_key).ratio()
                     for token_key in token_keys),
                    default=0,
                )
                if best_ratio >= 0.92:
                    fuzzy_count += 1
            match_count = exact_count + fuzzy_count
            if match_count:
                matches.append((session_id, match_count))

        if not matches:
            return None, 0, False

        matches.sort(key=lambda item: item[1], reverse=True)
        best_session_id, best_count = matches[0]
        ambiguous = len(matches) > 1 and matches[1][1] == best_count
        return best_session_id, best_count, ambiguous

    def delete_session(self, session_id: str):
        if session_id in self.mappings:
            del self.mappings[session_id]
            self._save_vault()

    def shred_vault(self):
        if self.vault_path.exists():
            self.vault_path.unlink()


class PresidioEngine:
    """Enhanced Presidio engine with custom recognizers."""

    def __init__(self, log_callback):
        self.log_callback = log_callback
        self.analyzer = None
        self._initialize_engine()

    def _initialize_engine(self):
        self.log_callback("🔧 Initializing Presidio Analyzer Engine...")
        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}],
        }
        try:
            provider = NlpEngineProvider(nlp_configuration=configuration)
            nlp_engine = provider.create_engine()
            self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
            self.log_callback("✅ SpaCy en_core_web_lg loaded")
        except Exception:
            self.log_callback("⚠️ en_core_web_lg not found, falling back to en_core_web_sm")
            configuration["models"][0]["model_name"] = "en_core_web_sm"
            provider = NlpEngineProvider(nlp_configuration=configuration)
            nlp_engine = provider.create_engine()
            self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine)

        self._add_insurance_policy_recognizer()
        self._add_medical_record_recognizer()
        self._add_drivers_license_recognizer()
        self._add_health_plan_recognizer()
        self._add_routing_number_recognizer()
        self._add_date_of_birth_recognizer()
        self._add_full_address_recognizer()
        self._add_street_address_recognizer()
        self._add_npi_recognizer()
        self._add_ssn_with_label_recognizer()
        self._add_financial_account_recognizer()
        self._add_cvv_recognizer()
        self._add_swift_code_recognizer()
        self._add_zip_code_recognizer()
        self._add_client_id_recognizer()
        self._add_loan_account_recognizer()
        self._add_generic_account_recognizer()

        self.log_callback("✅ Presidio Engine Ready with Custom Recognizers")

    # ------------------------------------------------------------------ #
    #  Custom recognizers                                                  #
    # ------------------------------------------------------------------ #

    def _add_insurance_policy_recognizer(self):
        p = Pattern("insurance_policy_pattern", r"\b(?:POL|POLICY|INS|P)-?\d{6,12}\b", 0.85)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="INSURANCE_POLICY", patterns=[p]))
        self.log_callback("  ➕ Insurance Policy Number recognizer added")

    def _add_medical_record_recognizer(self):
        p = Pattern("mrn_pattern",
                    r"\b(?:MRN|MR|MEDICAL.?RECORD|PATIENT.?ID)[:\s#-]*(\d{6,10})\b", 0.9)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="MEDICAL_RECORD_NUMBER", patterns=[p]))
        self.log_callback("  ➕ Medical Record Number (MRN) recognizer added")

    def _add_drivers_license_recognizer(self):
        p = Pattern("drivers_license_pattern",
                    r"\b(?:DL|DRIVER.?LICENSE)[:\s#-]*([A-Z0-9]{6,12})\b", 0.85)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="DRIVERS_LICENSE", patterns=[p]))
        self.log_callback("  ➕ Driver's License recognizer added")

    def _add_health_plan_recognizer(self):
        p = Pattern("health_plan_pattern",
                    r"\b(?:HEALTH.?PLAN|INSURANCE.?ID|MEMBER.?ID)[:\s#-]*([A-Z0-9]{8,15})\b", 0.85)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="HEALTH_PLAN_ID", patterns=[p]))
        self.log_callback("  ➕ Health Plan ID recognizer added")

    def _add_routing_number_recognizer(self):
        p = Pattern("routing_number_pattern",
                    r"\b(?:Routing(?:\s+Number)?|RTN|ABA)[:\s#-]*(\d{9})\b", 0.90)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="ROUTING_NUMBER", patterns=[p]))
        self.log_callback("  ➕ Routing Number recognizer added")

    def _add_date_of_birth_recognizer(self):
        p = Pattern("dob_pattern",
                    r"(?i)\b(?:DOB|Date\s+of\s+Birth|Birth\s+Date|Birthdate)"
                    r"[:\s]+\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}\b", 0.95)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="DATE_OF_BIRTH", patterns=[p]))
        self.log_callback("  ➕ Date of Birth recognizer added")

    def _add_full_address_recognizer(self):
        p = Pattern(
            "full_address_pattern",
            r"\b\d{1,6}[ \t]+[A-Z0-9][A-Za-z0-9\.\-]*"
            r"(?:[ \t]+[A-Z0-9][A-Za-z0-9\.\-]*){0,6}[ \t]+"
            r"(?:Street|St|Avenue|Ave|Boulevard|Blvd|Drive|Dr|Road|Rd"
            r"|Lane|Ln|Court|Ct|Way|Place|Pl|Parkway|Pkwy|Circle|Cir"
            r"|Trail|Trl|Highway|Hwy)\b\.?"
            r"(?:,[ \t]*(?:Suite|Ste|Apt|Unit)[ \t]*[A-Z0-9\-]+)?"
            r",[ \t]*[A-Z][A-Za-z]*(?:[ \t]+[A-Z][A-Za-z]*){0,3}"
            r",[ \t]*[A-Z]{2}[ \t]+\d{5}(?:-\d{4})?\b",
            0.96
        )
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="FULL_ADDRESS", patterns=[p]))
        self.log_callback("  + Full Address recognizer added")

    def _add_street_address_recognizer(self):
        p = Pattern(
            "street_address_pattern",
            r"\b\d{1,6}[ \t]+[A-Z0-9][A-Za-z0-9\.\-]*"
            r"(?:[ \t]+[A-Z0-9][A-Za-z0-9\.\-]*){0,6}[ \t]+"
            r"(?:Street|St|Avenue|Ave|Boulevard|Blvd|Drive|Dr|Road|Rd"
            r"|Lane|Ln|Court|Ct|Way|Place|Pl|Parkway|Pkwy|Circle|Cir"
            r"|Trail|Trl|Highway|Hwy)\b\.?"
            r"(?:,[ \t]*(?:Suite|Ste|Apt|Unit)[ \t]*[A-Z0-9\-]+)?",
            0.78
        )
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="STREET_ADDRESS", patterns=[p]))
        self.log_callback("  ➕ Street Address recognizer added")

    def _add_npi_recognizer(self):
        p = Pattern("npi_pattern",
                    r"\b(?:NPI|National\s+Provider)[:\s#]*(\d{10})\b", 0.9)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="NPI_NUMBER", patterns=[p]))
        self.log_callback("  ➕ NPI Number recognizer added")

    def _add_ssn_with_label_recognizer(self):
        """
        Catch SSNs labelled with 'SSN:' even when the label and value are
        separated by a newline (common in PDF text extraction).
        """
        p = Pattern(
            "ssn_with_label",
            r"(?i)(?:SSN|S\.S\.N|Social\s+Security(?:\s+Number)?)"
            r"[:\s#\n\r]*(\b(?!000|666|9\d{2})\d{3}[- ](?!00)\d{2}[- ](?!0000)\d{4}\b)",
            0.97
        )
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="US_SSN", patterns=[p]))
        self.log_callback("  ➕ SSN-with-label recognizer added")

    def _add_financial_account_recognizer(self):
        """
        Catch 16-digit account numbers (XXXX-XXXX-XXXX-XXXX) regardless of
        Luhn validity.  Also catches 15-digit Amex and 12-digit formats.
        """
        patterns = [
            Pattern("account_16digit", r"\b\d{4}[- ]\d{4}[- ]\d{4}[- ]\d{4}\b", 0.82),
            Pattern("account_15digit_amex", r"\b\d{4}[- ]\d{6}[- ]\d{5}\b", 0.82),
            Pattern("account_12digit", r"\b\d{4}[- ]\d{4}[- ]\d{4}\b", 0.76),
        ]
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="FINANCIAL_ACCOUNT", patterns=patterns))
        self.log_callback("  ➕ Financial Account Number recognizer added (16/15/12-digit)")

    def _add_cvv_recognizer(self):
        p = Pattern("cvv_pattern", r"(?i)\bCVV[:\s]*(\d{3,4})\b", 0.95)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="CVV_CODE", patterns=[p]))
        self.log_callback("  ➕ CVV Code recognizer added")

    def _add_swift_code_recognizer(self):
        p = Pattern(
            "swift_code_pattern",
            r"(?i)(?:Swift\s+Code|SWIFT|BIC)[:\s]*([A-Z]{4}[A-Z]{2}[A-Z0-9]{2}(?:[A-Z0-9]{3})?)\b",
            0.92
        )
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="SWIFT_CODE", patterns=[p]))
        self.log_callback("  ➕ SWIFT/BIC Code recognizer added")

    def _add_zip_code_recognizer(self):
        patterns = [
            # "NY 10022" or "IL 60611-1234"
            Pattern("zip_with_state", r"\b[A-Z]{2}\s+\d{5}(?:-\d{4})?\b", 0.78),
            # "Zip: 90210" or "ZIP Code: 10001"
            Pattern("zip_label", r"(?i)\bZip(?:\s+Code)?[:\s]+(\d{5}(?:-\d{4})?)\b", 0.92),
        ]
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="US_ZIP_CODE", patterns=patterns))
        self.log_callback("  ➕ US ZIP Code recognizer added")

    def _add_client_id_recognizer(self):
        """CLI-XXXX and CLI-JNT-XXXX patterns used in financial docs."""
        p = Pattern("client_id_pattern", r"\bCLI(?:-JNT)?-[A-Z0-9][A-Z0-9\-]{3,24}\b", 0.93)
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="CLIENT_ID", patterns=[p]))
        self.log_callback("  ➕ Client ID recognizer added")

    def _add_loan_account_recognizer(self):
        patterns = [
            Pattern("loan_account", r"\b(?:MTG|HELOC|LOC|LOAN)-[A-Z0-9][A-Z0-9\-]{3,24}\b", 0.92),
            Pattern("education_savings", r"\b529-[A-Z]{2}-[A-Z0-9][A-Z0-9\-]{3,18}\b", 0.92),
        ]
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="LOAN_ACCOUNT", patterns=patterns))
        self.log_callback("  ➕ Loan/Mortgage/529 Account recognizer added")

    def _add_generic_account_recognizer(self):
        """Catch account numbers that appear after explicit labels."""
        p = Pattern(
            "generic_account_labeled",
            r"\b(?:Account\s+(?:No|Number|#)|Acct\.?\s*(?:No|#)"
            r"|Patient\s+(?:No|Number|#)|Claim\s+(?:No|Number|#)"
            r"|Policy\s+(?:No|Number|#))[:\s]*([A-Z0-9\-]{5,20})\b",
            0.85
        )
        self.analyzer.registry.add_recognizer(
            PatternRecognizer(supported_entity="ACCOUNT_NUMBER", patterns=[p]))
        self.log_callback("  ➕ Generic labeled Account Number recognizer added")

    # ------------------------------------------------------------------ #
    #  Analysis                                                            #
    # ------------------------------------------------------------------ #

    def analyze_text(self, text: str) -> List:
        """Analyze text and return deduplicated, sorted PII/PHI entities."""
        entities = [
            # Core identity / contact
            "PERSON",
            "PHONE_NUMBER",
            "EMAIL_ADDRESS",
            "US_SSN",
            "US_ITIN",
            "US_PASSPORT",
            "US_DRIVER_LICENSE",
            "MEDICAL_LICENSE",
            "IP_ADDRESS",
            # Financial
            "CREDIT_CARD",
            "US_BANK_NUMBER",
            "IBAN_CODE",
            "FINANCIAL_ACCOUNT",
            "ROUTING_NUMBER",
            "SWIFT_CODE",
            "CVV_CODE",
            "CLIENT_ID",
            "LOAN_ACCOUNT",
            "ACCOUNT_NUMBER",
            # Medical
            "INSURANCE_POLICY",
            "MEDICAL_RECORD_NUMBER",
            "DRIVERS_LICENSE",
            "HEALTH_PLAN_ID",
            "NPI_NUMBER",
            # Address / location (kept but requires higher confidence via threshold)
            "FULL_ADDRESS",
            "STREET_ADDRESS",
            "US_ZIP_CODE",
            # DATE_OF_BIRTH — specific DOB detector only; generic DATE_TIME removed
            # (transaction dates, filing dates, etc. are not PII)
            "DATE_OF_BIRTH",
            # NRP removed — fires on nationalities, religions, political affiliations
            # in normal document prose and causes many false positives
        ]

        results = self.analyzer.analyze(
            text=text,
            entities=entities,
            language='en',
            score_threshold=0.5,   # raised from 0.3 — cuts low-confidence noise
            allow_list=None,
        )

        # Remove overlapping spans — keep highest-confidence detection per span
        results = self._deduplicate_results(results)
        # Reverse-sort for safe right-to-left position processing (used as reference only)
        results.sort(key=lambda x: x.start, reverse=True)
        return results

    @staticmethod
    def _deduplicate_results(results: List) -> List:
        """Keep only the highest-scoring result when two spans overlap."""
        sorted_results = sorted(results, key=lambda x: x.score, reverse=True)
        kept = []
        for candidate in sorted_results:
            overlaps = any(
                candidate.start < k.end and candidate.end > k.start
                for k in kept
            )
            if not overlaps:
                kept.append(candidate)
        return kept


class PIISanitizerApp(TkinterDnD.Tk):
    """Main application window with drag-and-drop support."""

    def __init__(self):
        super().__init__()

        self.vault = SessionVault()

        self.title("Peoples First Insurance - Privacy Sanitizer")
        self.geometry("1100x760")
        self.configure(bg="#f4f7fb")

        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        self.presidio_engine = None
        self.engine_ready = False
        self.mask_option_vars = {}

        self.setup_ui()
        self.initialize_engine()
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

    def setup_ui(self):
        """Create the user interface."""
        # Header
        self.header_frame = ctk.CTkFrame(self, fg_color="#ffffff", corner_radius=0)
        self.header_frame.pack(fill="x", padx=0, pady=0)

        self.brand_frame = ctk.CTkFrame(self.header_frame, fg_color="#ffffff")
        brand_frame = self.brand_frame
        brand_frame.pack(side="left", padx=24, pady=10)
        self._create_brand_logo(brand_frame).pack(side="left")

        self.product_frame = ctk.CTkFrame(self.header_frame, fg_color="#ffffff")
        product_frame = self.product_frame
        product_frame.pack(side="right", padx=(0, 24), pady=12)

        self.product_title_label = ctk.CTkLabel(
            product_frame,
            text="Privacy Sanitizer",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color="#17365d",
            anchor="e"
        )
        self.product_title_label.pack(anchor="e")

        self.product_subtitle_label = ctk.CTkLabel(
            product_frame,
            text="Tokenize, restore, and mask protected data",
            font=ctk.CTkFont(size=12),
            text_color="#667085",
            anchor="e"
        )
        self.product_subtitle_label.pack(anchor="e", pady=(2, 0))

        self.theme_switch = ctk.CTkSwitch(
            product_frame,
            text="Dark mode",
            command=self.toggle_theme,
            font=ctk.CTkFont(size=12),
            text_color="#1f2937",
            progress_color="#8a1538"
        )
        self.theme_switch.pack(anchor="e", pady=(8, 0))

        self.accent_bar = ctk.CTkFrame(self, fg_color="#8a1538", height=4, corner_radius=0)
        self.accent_bar.pack(fill="x")

        # Main content
        self.content_frame = ctk.CTkFrame(self, fg_color="transparent")
        content_frame = self.content_frame
        content_frame.pack(fill="both", expand=True, padx=24, pady=22)

        self.instructions_label = ctk.CTkLabel(
            content_frame,
            text="Tokenize documents for reversible AI workflows, restore tokenized files, or permanently mask selected data.",
            font=ctk.CTkFont(size=13),
            text_color="#aaaaaa"
        )
        self.instructions_label.pack(fill="x", pady=(0, 15))

        workflow_tabs = ctk.CTkTabview(
            content_frame,
            fg_color="#ffffff",
            segmented_button_selected_color="#d7e5f0",
            segmented_button_selected_hover_color="#c4d9eb",
            segmented_button_unselected_color="#ffffff",
            segmented_button_unselected_hover_color="#edf3f8",
            text_color="#17365d",
        )
        self.workflow_tabs = workflow_tabs
        workflow_tabs.pack(fill="both", expand=True, pady=(0, 20))
        token_tab = workflow_tabs.add("Tokenize & Restore")
        mask_tab = workflow_tabs.add("Mask")

        # Two-column layout
        columns_frame = ctk.CTkFrame(token_tab, fg_color="transparent")
        columns_frame.pack(fill="both", expand=True, padx=12, pady=12)

        # --- Left: SANITIZE ---
        left_column = ctk.CTkFrame(columns_frame, fg_color="transparent")
        left_column.pack(side="left", fill="both", expand=True, padx=(0, 10))

        self.sanitize_title_label = ctk.CTkLabel(
            left_column,
            text="Sanitize Document",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color="#3b8ed0"
        )
        self.sanitize_title_label.pack(pady=(0, 10))

        self.sanitize_drop_frame = ctk.CTkFrame(
            left_column, fg_color="#ffffff",
            border_width=2, border_color="#1f5f99",
            corner_radius=8, height=210
        )
        self.sanitize_drop_frame.pack(fill="both", expand=True, pady=(0, 15))
        self.sanitize_drop_frame.pack_propagate(False)

        self.sanitize_drop_label = ctk.CTkLabel(
            self.sanitize_drop_frame,
            text="Drop source file here\n\nTokenize or mask sensitive data",
            font=ctk.CTkFont(size=14), text_color="#aaaaaa"
        )
        self.sanitize_drop_label.place(relx=0.5, rely=0.5, anchor="center")

        self.sanitize_drop_frame.drop_target_register(DND_FILES)
        self.sanitize_drop_frame.dnd_bind('<<Drop>>', self.on_sanitize_drop)

        self.sanitize_btn = ctk.CTkButton(
            left_column, text="SANITIZE FILE",
            font=ctk.CTkFont(size=14, weight="bold"), height=45,
            command=self.sanitize_current_file, state="disabled",
            fg_color="#1f5f99", hover_color="#17365d", corner_radius=6
        )
        self.sanitize_btn.pack(fill="x")

        # --- Right: REHYDRATE ---
        right_column = ctk.CTkFrame(columns_frame, fg_color="transparent")
        right_column.pack(side="right", fill="both", expand=True, padx=(10, 0))

        self.restore_title_label = ctk.CTkLabel(
            right_column,
            text="Restore Tokenized Document",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color="#2c8c2c"
        )
        self.restore_title_label.pack(pady=(0, 10))

        self.rehydrate_drop_frame = ctk.CTkFrame(
            right_column, fg_color="#ffffff",
            border_width=2, border_color="#247a4d",
            corner_radius=8, height=210
        )
        self.rehydrate_drop_frame.pack(fill="both", expand=True, pady=(0, 15))
        self.rehydrate_drop_frame.pack_propagate(False)

        self.rehydrate_drop_label = ctk.CTkLabel(
            self.rehydrate_drop_frame,
            text="Drop tokenized file here\n\nUses session or token matching",
            font=ctk.CTkFont(size=14), text_color="#aaaaaa"
        )
        self.rehydrate_drop_label.place(relx=0.5, rely=0.5, anchor="center")

        self.rehydrate_drop_frame.drop_target_register(DND_FILES)
        self.rehydrate_drop_frame.dnd_bind('<<Drop>>', self.on_rehydrate_drop)

        self.rehydrate_btn = ctk.CTkButton(
            right_column, text="RESTORE FILE",
            font=ctk.CTkFont(size=14, weight="bold"), height=45,
            command=self.rehydrate_current_file, state="disabled",
            fg_color="#247a4d", hover_color="#1e5e3c", corner_radius=6
        )
        self.rehydrate_btn.pack(fill="x")

        # --- Mask tab ---
        mask_layout = ctk.CTkFrame(mask_tab, fg_color="transparent")
        mask_layout.pack(fill="both", expand=True, padx=12, pady=12)

        self.mask_options_frame = ctk.CTkFrame(mask_layout, fg_color="#f8fafc", corner_radius=8)
        options_frame = self.mask_options_frame
        options_frame.pack(side="left", fill="y", padx=(0, 12))

        self.mask_options_title_label = ctk.CTkLabel(
            options_frame,
            text="Masking Options",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color="#17365d",
            anchor="w"
        )
        self.mask_options_title_label.pack(fill="x", padx=16, pady=(16, 4))

        self.mask_options_help_label = ctk.CTkLabel(
            options_frame,
            text="Masking is permanent and cannot be restored.",
            font=ctk.CTkFont(size=12),
            text_color="#667085",
            anchor="w",
            wraplength=250
        )
        self.mask_options_help_label.pack(fill="x", padx=16, pady=(0, 12))

        mask_options = [
            ("All", "all"),
            ("Names and identity", "identity"),
            ("Email addresses", "email"),
            ("Phone numbers", "phone"),
            ("IP addresses", "ip"),
            ("Financial data", "financial"),
            ("Addresses and ZIP codes", "address"),
            ("Medical and insurance IDs", "medical"),
            ("Dates of birth", "dates"),
        ]
        self.mask_option_widgets = []
        for label, key in mask_options:
            var = ctk.BooleanVar(value=True if key == "all" else False)
            self.mask_option_vars[key] = var
            command = self.on_mask_all_toggle if key == "all" else self.on_mask_option_toggle
            checkbox = ctk.CTkCheckBox(
                options_frame,
                text=label,
                variable=var,
                command=command,
                text_color="#1f2937",
                fg_color="#1f5f99",
                hover_color="#17365d",
            )
            checkbox.pack(anchor="w", padx=16, pady=6)
            self.mask_option_widgets.append(checkbox)

        mask_work_frame = ctk.CTkFrame(mask_layout, fg_color="transparent")
        mask_work_frame.pack(side="right", fill="both", expand=True)

        self.mask_title_label = ctk.CTkLabel(
            mask_work_frame,
            text="Permanently Mask Document",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color="#c2410c"
        )
        self.mask_title_label.pack(pady=(0, 10))

        self.mask_drop_frame = ctk.CTkFrame(
            mask_work_frame, fg_color="#ffffff",
            border_width=2, border_color="#c2410c",
            corner_radius=8, height=210
        )
        self.mask_drop_frame.pack(fill="both", expand=True, pady=(0, 15))
        self.mask_drop_frame.pack_propagate(False)

        self.mask_drop_label = ctk.CTkLabel(
            self.mask_drop_frame,
            text="Drop source file here\n\nSelected fields will be irreversibly masked",
            font=ctk.CTkFont(size=14), text_color="#aaaaaa"
        )
        self.mask_drop_label.place(relx=0.5, rely=0.5, anchor="center")

        self.mask_drop_frame.drop_target_register(DND_FILES)
        self.mask_drop_frame.dnd_bind('<<Drop>>', self.on_mask_drop)

        self.mask_btn = ctk.CTkButton(
            mask_work_frame, text="MASK FILE",
            font=ctk.CTkFont(size=14, weight="bold"), height=45,
            command=self.mask_current_file, state="disabled",
            fg_color="#c2410c", hover_color="#9a3412", corner_radius=6
        )
        self.mask_btn.pack(fill="x")

        # Activity feed
        activity_header = ctk.CTkFrame(content_frame, fg_color="transparent")
        activity_header.pack(fill="x", pady=(0, 5))

        self.activity_title_label = ctk.CTkLabel(
            activity_header, text="Activity",
            font=ctk.CTkFont(size=14, weight="bold"), anchor="w",
            text_color="#1f2937"
        )
        self.activity_title_label.pack(side="left")

        self.activity_toggle = ctk.CTkSwitch(
            activity_header,
            text="Show activity",
            command=self.toggle_activity_feed,
            font=ctk.CTkFont(size=12),
            text_color="#667085",
            progress_color="#1f5f99",
        )
        self.activity_toggle.select()
        self.activity_toggle.pack(side="right")

        self.activity_container = ctk.CTkScrollableFrame(
            content_frame,
            fg_color="#ffffff",
            height=160,
            corner_radius=8,
        )
        self.activity_container.pack(fill="both", expand=False)
        self.activity_items = []

        # Status bar
        self.status_label = ctk.CTkLabel(
            self, text="⏳ Initializing Presidio Engine...",
            font=ctk.CTkFont(size=11), anchor="w",
            fg_color="#17365d", text_color="#ffffff", corner_radius=0
        )
        self.status_label.pack(fill="x", padx=0, pady=0)

        self.current_sanitize_file = None
        self.current_rehydrate_file = None
        self.current_mask_file = None

        self.log("🚀 Application started. Waiting for Presidio initialization...")
        if not PYMUPDF_AVAILABLE and not PDF_AVAILABLE:
            self.log("⚠️  No PDF library found — PDF input will not work.")
            self.log("   Run:  pip install pymupdf   then restart the app.")

    def _create_brand_logo(self, parent):
        canvas = tk.Canvas(
            parent,
            width=520,
            height=96,
            bg="#ffffff",
            highlightthickness=0,
            bd=0,
        )

        burgundy = "#8a1538"
        black = "#1f1f1f"
        gray = "#2d2d2d"

        canvas.create_arc(
            8, 2, 118, 82,
            start=112,
            extent=250,
            style="arc",
            outline=burgundy,
            width=9,
        )
        canvas.create_text(
            55, 54,
            text="P",
            fill=black,
            font=("Times New Roman", 74, "italic"),
        )
        canvas.create_text(
            135, 34,
            text="PEOPLES",
            fill=black,
            anchor="w",
            font=("Georgia", 33, "bold"),
        )
        canvas.create_text(
            356, 34,
            text="FIRST",
            fill=burgundy,
            anchor="w",
            font=("Georgia", 33, "bold"),
        )
        canvas.create_text(
            137, 75,
            text="I N S U R A N C E",
            fill=gray,
            anchor="w",
            font=("Georgia", 23),
        )

        return canvas

    def initialize_engine(self):
        def init_thread():
            self.presidio_engine = PresidioEngine(self.log)
            self.engine_ready = True
            self.after(0, self._enable_buttons)
        threading.Thread(target=init_thread, daemon=True).start()

    def _enable_buttons(self):
        self.status_label.configure(text="✅ Ready - Drop files to begin")
        self.log("🟢 System Ready. Drop files into the zones above to begin.")

    def on_mask_all_toggle(self):
        all_enabled = self.mask_option_vars["all"].get()
        for key, var in self.mask_option_vars.items():
            if key != "all":
                var.set(False if all_enabled else var.get())

    def on_mask_option_toggle(self):
        any_specific = any(
            var.get() for key, var in self.mask_option_vars.items() if key != "all"
        )
        self.mask_option_vars["all"].set(not any_specific)

    def get_selected_mask_entities(self) -> Optional[Set[str]]:
        if self.mask_option_vars.get("all") and self.mask_option_vars["all"].get():
            return None

        groups = {
            "identity": {
                "PERSON", "US_SSN", "US_ITIN", "US_PASSPORT",
                "US_DRIVER_LICENSE", "DRIVERS_LICENSE", "CLIENT_ID",
            },
            "email": {"EMAIL_ADDRESS"},
            "phone": {"PHONE_NUMBER"},
            "ip": {"IP_ADDRESS"},
            "financial": {
                "CREDIT_CARD", "US_BANK_NUMBER", "IBAN_CODE",
                "FINANCIAL_ACCOUNT", "ROUTING_NUMBER", "SWIFT_CODE",
                "CVV_CODE", "LOAN_ACCOUNT", "ACCOUNT_NUMBER",
            },
            "address": {"FULL_ADDRESS", "STREET_ADDRESS", "US_ZIP_CODE"},
            "medical": {
                "INSURANCE_POLICY", "MEDICAL_RECORD_NUMBER", "MEDICAL_LICENSE",
                "HEALTH_PLAN_ID", "NPI_NUMBER",
            },
            "dates": {"DATE_OF_BIRTH"},
        }
        selected = set()
        for key, entities in groups.items():
            if self.mask_option_vars.get(key) and self.mask_option_vars[key].get():
                selected.update(entities)
        return selected

    def toggle_theme(self):
        if self.theme_switch.get():
            ctk.set_appearance_mode("dark")
            self._apply_theme("dark")
            self.log("Switched to dark mode")
        else:
            ctk.set_appearance_mode("light")
            self._apply_theme("light")
            self.log("Switched to light mode")
        self.update_idletasks()

    def toggle_activity_feed(self):
        if self.activity_toggle.get():
            self.activity_container.pack(fill="both", expand=False)
        else:
            self.activity_container.pack_forget()

    def _theme_palette(self, mode: str) -> dict:
        if mode == "dark":
            return {
                "app_bg": "#111827",
                "header": "#1e293b",
                "header_text": "#f9fafb",
                "drop": "#253244",
                "panel_alt": "#172033",
                "text": "#f9fafb",
                "muted": "#cbd5e1",
                "tab": "#1f2937",
                "tab_selected": "#1f5f99",
                "tab_unselected": "#334155",
                "tab_hover": "#475569",
                "activity": "#1f2937",
                "activity_item": "#253244",
            }
        return {
            "app_bg": "#f4f7fb",
            "header": "#ffffff",
            "header_text": "#17365d",
            "drop": "#ffffff",
            "panel_alt": "#f8fafc",
            "text": "#1f2937",
            "muted": "#667085",
            "tab": "#ffffff",
            "tab_selected": "#d7e5f0",
            "tab_unselected": "#ffffff",
            "tab_hover": "#edf3f8",
            "activity": "#ffffff",
            "activity_item": "#f8fafc",
        }

    def _apply_theme(self, mode: str):
        colors = self._theme_palette(mode)

        # Root window
        self.configure(bg=colors["app_bg"])

        # Header bar + product area
        self.header_frame.configure(fg_color=colors["header"])
        self.brand_frame.configure(fg_color=colors["header"])
        self.product_frame.configure(fg_color=colors["header"])
        self.product_title_label.configure(text_color=colors["header_text"])
        self.product_subtitle_label.configure(text_color=colors["muted"])
        self.theme_switch.configure(text_color=colors["header_text"])

        # Main content
        self.instructions_label.configure(text_color=colors["muted"])
        self.workflow_tabs.configure(
            fg_color=colors["tab"],
            segmented_button_selected_color=colors["tab_selected"],
            segmented_button_selected_hover_color="#c4d9eb",
            segmented_button_unselected_color=colors["tab_unselected"],
            segmented_button_unselected_hover_color=colors["tab_hover"],
            text_color=colors["header_text"],
        )
        for frame in [self.sanitize_drop_frame, self.rehydrate_drop_frame, self.mask_drop_frame]:
            frame.configure(fg_color=colors["drop"])
        self.mask_options_frame.configure(fg_color=colors["panel_alt"])

        # Activity area — configure the outer frame then force the inner
        # tk.Canvas (CTkScrollableFrame uses a raw canvas for scrolling that
        # does NOT automatically repaint when the CTk fg_color is set).
        self.activity_container.configure(fg_color=colors["activity"])
        try:
            self.activity_container._parent_canvas.configure(bg=colors["activity"])
        except Exception:
            pass

        for label in [
            self.sanitize_drop_label, self.rehydrate_drop_label,
            self.mask_drop_label, self.mask_options_help_label,
            self.activity_toggle,
        ]:
            label.configure(text_color=colors["muted"])
        for label in [self.activity_title_label, self.mask_options_title_label]:
            label.configure(text_color=colors["text"])
        for checkbox in getattr(self, "mask_option_widgets", []):
            checkbox.configure(text_color=colors["text"])
        for item in getattr(self, "activity_items", []):
            item["frame"].configure(fg_color=colors["activity_item"])
            item["label"].configure(text_color=colors["text"])

    def log(self, message: str):
        timestamp = datetime.now().strftime("%H:%M:%S")
        if not hasattr(self, "activity_container"):
            return
        colors = self._theme_palette("dark" if self.theme_switch.get() else "light")
        item = ctk.CTkFrame(self.activity_container, fg_color=colors["activity_item"], corner_radius=6)
        item.pack(fill="x", padx=8, pady=(6, 0))
        label = ctk.CTkLabel(
            item,
            text=f"{timestamp}  {message}",
            font=ctk.CTkFont(size=12),
            text_color=colors["text"],
            anchor="w",
            justify="left",
            wraplength=980,
        )
        label.pack(fill="x", padx=10, pady=7)
        self.activity_items.append({"frame": item, "label": label})
        if len(self.activity_items) > 80:
            old = self.activity_items.pop(0)
            old["frame"].destroy()
        self.activity_container._parent_canvas.yview_moveto(1.0)
        self.update_idletasks()

    def on_sanitize_drop(self, event):
        if not self.engine_ready:
            self.log("⚠️ Please wait for engine initialization to complete.")
            return
        file_path = event.data.strip('{}')
        self.current_sanitize_file = file_path
        self.sanitize_drop_label.configure(
            text=f"✅ {Path(file_path).name}\n\nReady to sanitize",
            text_color="#00ff00"
        )
        self.sanitize_btn.configure(state="normal")
        self.log(f"📥 File ready for sanitization: {Path(file_path).name}")

    def on_rehydrate_drop(self, event):
        if not self.engine_ready:
            self.log("⚠️ Please wait for engine initialization to complete.")
            return
        file_path = event.data.strip('{}')
        self.current_rehydrate_file = file_path
        self.rehydrate_drop_label.configure(
            text=f"✅ {Path(file_path).name}\n\nReady to rehydrate",
            text_color="#00ff00"
        )
        self.rehydrate_btn.configure(state="normal")
        self.log(f"📥 File ready for rehydration: {Path(file_path).name}")

    def on_mask_drop(self, event):
        if not self.engine_ready:
            self.log("Please wait for engine initialization to complete.")
            return
        file_path = event.data.strip('{}')
        self.current_mask_file = file_path
        self.mask_drop_label.configure(
            text=f"✅ {Path(file_path).name}\n\nReady to mask",
            text_color="#16a34a"
        )
        self.mask_btn.configure(state="normal")
        self.log(f"File ready for masking: {Path(file_path).name}")

    def sanitize_current_file(self):
        if self.current_sanitize_file:
            self.sanitize_file(self.current_sanitize_file, mode="tokenize")
            self.current_sanitize_file = None
            self.sanitize_drop_label.configure(
                text="Drop source file here\n\nCreate reversible privacy tokens",
                text_color="#aaaaaa"
            )
            self.sanitize_btn.configure(state="disabled")

    def mask_current_file(self):
        if self.current_mask_file:
            self.sanitize_file(
                self.current_mask_file,
                mode="mask",
                mask_entities=self.get_selected_mask_entities(),
            )
            self.current_mask_file = None
            self.mask_drop_label.configure(
                text="Drop source file here\n\nSelected fields will be irreversibly masked",
                text_color="#aaaaaa"
            )
            self.mask_btn.configure(state="disabled")

    def rehydrate_current_file(self):
        if self.current_rehydrate_file:
            self.rehydrate_file(self.current_rehydrate_file)
            self.current_rehydrate_file = None
            self.rehydrate_drop_label.configure(
                text="Drop tokenized file here\n\nUses session or token matching",
                text_color="#aaaaaa"
            )
            self.rehydrate_btn.configure(state="disabled")

    # ------------------------------------------------------------------ #
    #  Core: sanitize                                                      #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _clean_detected_value(entity_type: str, value: str) -> str:
        """Keep labels in the document and tokenize only the sensitive value."""
        value = value.strip()
        patterns = {
            "US_SSN": r"\b(?!000|666|9\d{2})\d{3}[- ](?!00)\d{2}[- ](?!0000)\d{4}\b",
            "ROUTING_NUMBER": r"\b\d{9}\b",
            "SWIFT_CODE": r"\b[A-Z]{4}[A-Z]{2}[A-Z0-9]{2}(?:[A-Z0-9]{3})?\b",
            "DATE_OF_BIRTH": r"\b\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}\b",
        }

        pattern = patterns.get(entity_type)
        if pattern:
            matches = re.findall(pattern, value, flags=re.IGNORECASE)
            if matches:
                return matches[-1] if isinstance(matches[-1], str) else matches[-1][0]

        if entity_type == "ACCOUNT_NUMBER":
            tail = re.split(r"[:\n\r]+", value, maxsplit=1)[-1].strip()
            match = re.search(r"\b[A-Z0-9][A-Z0-9\-]{4,24}\b", tail, flags=re.IGNORECASE)
            if match:
                return match.group(0)

        return value

    @staticmethod
    def _last_digits(value: str, count: int = 4) -> str:
        digits = re.sub(r"\D", "", value or "")
        return digits[-count:] if len(digits) >= count else digits

    @staticmethod
    def _mask_like_original(value: str, visible_tail: str, mask_char: str = "x") -> str:
        tail_index = len(visible_tail) - 1
        result = []
        for char in reversed(value):
            if char.isdigit():
                if tail_index >= 0 and char == visible_tail[tail_index]:
                    result.append(char)
                    tail_index -= 1
                else:
                    result.append(mask_char)
            elif char.isalpha():
                result.append(mask_char)
            else:
                result.append(char)
        return "".join(reversed(result))

    @staticmethod
    def _mask_value(entity_type: str, value: str) -> str:
        value = value.strip()
        tail4 = PIISanitizerApp._last_digits(value, 4)

        if entity_type in {"US_SSN", "US_ITIN"} and tail4:
            return f"xxx-xx-{tail4}"
        if entity_type in {"PHONE_NUMBER"} and tail4:
            return f"(xxx) xxx-{tail4}"
        if entity_type in {
            "CREDIT_CARD", "FINANCIAL_ACCOUNT", "ACCOUNT_NUMBER",
            "US_BANK_NUMBER", "IBAN_CODE"
        } and tail4:
            return f"xxxx-xxxx-xxxx-{tail4}"
        if entity_type == "ROUTING_NUMBER" and tail4:
            return f"xxxxx{tail4}"
        if entity_type == "CVV_CODE":
            return "CVV: xxx" if value.lower().startswith("cvv") else "xxx"
        if entity_type == "EMAIL_ADDRESS" and "@" in value:
            _, domain = value.split("@", 1)
            return f"masked@{domain}"
        if entity_type in {"DATE_OF_BIRTH"}:
            year_match = re.search(r"(\d{4})\b", value)
            return f"xx/xx/{year_match.group(1)}" if year_match else "xx/xx/xxxx"
        if entity_type in {"FULL_ADDRESS", "STREET_ADDRESS"}:
            return "[masked address]"
        if entity_type == "US_ZIP_CODE":
            return "xxxxx"
        if entity_type == "PERSON":
            return "[masked name]"
        if entity_type in {"CLIENT_ID", "LOAN_ACCOUNT", "INSURANCE_POLICY"} and tail4:
            return f"masked-{tail4}"
        if entity_type in {
            "SWIFT_CODE", "US_PASSPORT", "US_DRIVER_LICENSE",
            "DRIVERS_LICENSE", "MEDICAL_LICENSE", "MEDICAL_RECORD_NUMBER",
            "HEALTH_PLAN_ID", "NPI_NUMBER", "IP_ADDRESS"
        }:
            return "[masked id]"

        if tail4:
            return PIISanitizerApp._mask_like_original(value, tail4)
        return "[masked]"

    def sanitize_file(self, file_path: str, mode: str = "tokenize",
                      mask_entities: Optional[Set[str]] = None):
        """Sanitize a file by replacing PII/PHI with tokens."""
        try:
            file_name = Path(file_path).name
            self.log(f"🔍 Scanning: {file_name}")
            self.status_label.configure(text="⏳ Analyzing file for PII/PHI...")

            try:
                original_text, file_format = FileHandler.read_file(file_path)
                self.log(f"📄 Detected format: {file_format.upper()}")
            except ImportError as e:
                self.log(f"❌ Missing library: {e}")
                self.status_label.configure(text="❌ Missing required library")
                return
            except Exception as e:
                self.log(f"❌ Failed to read file: {e}")
                self.status_label.configure(text="❌ File read error")
                return

            results = self.presidio_engine.analyze_text(original_text)
            if mode == "mask" and mask_entities is not None:
                results = [result for result in results if result.entity_type in mask_entities]

            if not results:
                self.log("✅ No PII/PHI detected. File is clean.")
                self.status_label.configure(text="✅ No sensitive data found")
                return

            irreversible_mask = mode == "mask"
            session_id = None if irreversible_mask else (
                f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                f"_{secrets.token_hex(4)}"
            )

            # --- Build value→token map (deduplicated by value) ---
            entity_counts: Dict[str, int] = {}
            value_to_token: Dict[str, str] = {}

            for result in results:
                entity_type = result.entity_type
                # Strip leading/trailing whitespace — Presidio sometimes includes
                # a surrounding space in the span, which would cause the PDF search
                # to match an extra space and leave a double-redacted gap.
                original_value = self._clean_detected_value(
                    entity_type,
                    original_text[result.start:result.end],
                )
                if len(original_value) < 2:   # ignore single-char or empty spans
                    continue
                if original_value not in value_to_token:
                    if irreversible_mask:
                        value_to_token[original_value] = self._mask_value(entity_type, original_value)
                    else:
                        token = self.vault.generate_token(entity_type)
                        self.vault.store_mapping(session_id, token, original_value)
                        value_to_token[original_value] = token
                entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1

            # --- Single-pass string replacement (no position-shift corruption) ---
            # Sort longest first so shorter substrings aren't replaced inside longer ones
            sanitized_text = original_text
            for original_value in sorted(value_to_token, key=len, reverse=True):
                sanitized_text = sanitized_text.replace(
                    original_value, value_to_token[original_value]
                )

            # --- Write output ---
            original_path = Path(file_path)
            output_suffix = "_MASKED" if irreversible_mask else "_SANITIZED"
            output_path = original_path.parent / f"{original_path.stem}{output_suffix}{original_path.suffix}"

            if file_format == 'pdf' and PYMUPDF_AVAILABLE:
                try:
                    count = (
                        FileHandler.write_pdf_masked(file_path, str(output_path), value_to_token)
                        if irreversible_mask else
                        FileHandler.write_pdf_sanitized(file_path, str(output_path), value_to_token)
                    )
                    action = "masked" if irreversible_mask else "tokenized"
                    self.log(f"PDF {action} in-place - {count} value(s) processed")
                except Exception as e:
                    self.log(f"⚠️ PDF write failed ({e}), saving as .txt fallback")
                    output_path = original_path.parent / f"{original_path.stem}{output_suffix}.txt"
                    FileHandler._write_text(output_path, sanitized_text)
            elif file_format == 'pdf':
                # No PyMuPDF — .txt is the only clean fallback
                output_path = original_path.parent / f"{original_path.stem}{output_suffix}.txt"
                FileHandler._write_text(output_path, sanitized_text)
                self.log("📄 Saved as .txt — run: pip install pymupdf  for PDF output")
            elif file_format == 'docx':
                try:
                    count = (
                        FileHandler.write_docx_masked(file_path, str(output_path), value_to_token)
                        if irreversible_mask else
                        FileHandler.write_docx_sanitized(file_path, str(output_path), value_to_token)
                    )
                    action = "masked" if irreversible_mask else "tokenized"
                    self.log(f"DOCX {action} in-place - {count} value(s) processed")
                except Exception as e:
                    self.log(f"DOCX write failed ({e}), saving as .txt fallback")
                    output_path = original_path.parent / f"{original_path.stem}{output_suffix}.txt"
                    FileHandler._write_text(output_path, sanitized_text)
            else:
                try:
                    FileHandler.write_file(str(output_path), sanitized_text, file_format)
                except Exception as e:
                    self.log(f"⚠️ Could not write as {file_format.upper()}, saving as .txt")
                    output_path = original_path.parent / f"{original_path.stem}{output_suffix}.txt"
                    FileHandler._write_text(output_path, sanitized_text)

            if not irreversible_mask:
                metadata_path = original_path.parent / f"{original_path.stem}_SANITIZED.session"
                with open(metadata_path, 'w') as f:
                    json.dump({
                        'session_id': session_id,
                        'original_format': file_format,
                        'original_extension': original_path.suffix
                    }, f, indent=2)

            self.log("✅ SANITIZATION COMPLETE")
            self.log("📊 Detected and Scrubbed:")
            for entity_type, count in sorted(entity_counts.items()):
                self.log(f"   • {count}x {entity_type}")
            self.log(f"💾 Safe file saved: {output_path.name}")
            if not irreversible_mask:
                self.log(f"Session ID: {session_id}")
            else:
                self.log("Mask mode is irreversible - no session was created")
            self.status_label.configure(
                text=f"✅ Scrubbed {sum(entity_counts.values())} sensitive items"
            )

        except Exception as e:
            self.log(f"❌ Error: {e}")
            self.status_label.configure(text="❌ Sanitization failed")
            import traceback
            self.log(f"Debug: {traceback.format_exc()}")

    # ------------------------------------------------------------------ #
    #  Core: rehydrate                                                     #
    # ------------------------------------------------------------------ #

    @staticmethod
    def _extract_tokens(text: str) -> Set[str]:
        text = text or ""
        tokens = set(re.findall(r"<[A-Z][A-Z0-9_]*_[A-Z0-9]{4}>", text))
        tokens.update(re.findall(r"\bTOKEN_\d{6}_[A-Z][A-Z0-9_]*\b", text))
        tokens.update(re.findall(r"\bToken_\d+\b", text, flags=re.IGNORECASE))
        return tokens

    @staticmethod
    def _extract_token_candidates(text: str) -> Set[str]:
        candidates = set(PIISanitizerApp._extract_tokens(text))
        for line in (text or "").splitlines():
            for match in re.finditer(r"\bToken[\s_\-]*\d+\b", line, flags=re.IGNORECASE):
                candidates.add(match.group(0).strip())
            for match in re.finditer(r"\bTOKEN[\s_\-]*\d{1,6}[\s_\-]+[A-Z][A-Z0-9_\s\-]{1,50}\b", line):
                candidates.add(match.group(0).strip())
            start = line.find("<")
            while start != -1:
                end = line.find(">", start + 1)
                if end == -1:
                    candidate = line[start:].strip()
                    next_start = -1
                else:
                    candidate = line[start:end + 1].strip()
                    next_start = line.find("<", end + 1)
                if 4 <= len(candidate) <= 90:
                    candidates.add(candidate)
                start = next_start
        return candidates

    @staticmethod
    def _build_rehydration_map(mappings: Dict[str, str], text: str) -> Dict[str, str]:
        replacement_map = dict(mappings)
        candidates = PIISanitizerApp._extract_token_candidates(text)
        canonical_by_key = {
            SessionVault.token_match_key(token): token
            for token in mappings
        }

        for candidate in candidates:
            if candidate in replacement_map:
                continue
            candidate_key = SessionVault.token_match_key(candidate)
            if not candidate_key:
                continue

            canonical = canonical_by_key.get(candidate_key)
            if canonical:
                replacement_map[candidate] = mappings[canonical]
                continue

            scored = []
            for token in mappings:
                token_key = SessionVault.token_match_key(token)
                ratio = difflib.SequenceMatcher(None, candidate_key, token_key).ratio()
                if ratio >= 0.92:
                    scored.append((ratio, token))
            scored.sort(reverse=True)
            if scored and (len(scored) == 1 or scored[0][0] - scored[1][0] >= 0.03):
                replacement_map[candidate] = mappings[scored[0][1]]

        return replacement_map

    @staticmethod
    def _session_file_points_to(session_file: Path, session_id: str) -> bool:
        try:
            content = session_file.read_text(encoding='utf-8').strip()
            if content.startswith('{'):
                return json.loads(content).get('session_id') == session_id
            return content == session_id
        except Exception:
            return False

    def _delete_session_files_for_session(
            self,
            directory: Path,
            session_id: str,
            primary_session_path: Optional[Path]) -> int:
        deleted = 0
        candidates = []
        if primary_session_path is not None:
            candidates.append(primary_session_path)
        candidates.extend(directory.glob("*.session"))

        seen = set()
        for session_file in candidates:
            if session_file in seen or not session_file.exists():
                continue
            seen.add(session_file)
            if self._session_file_points_to(session_file, session_id):
                session_file.unlink()
                deleted += 1
        return deleted

    def rehydrate_file(self, file_path: str):
        """Rehydrate a sanitized file by restoring original PII/PHI."""
        try:
            self.log(f"💧 Rehydrating: {Path(file_path).name}")
            self.status_label.configure(text="⏳ Restoring original data...")

            original_path = Path(file_path)
            session_path = original_path.parent / f"{original_path.stem}.session"
            session_source = "sidecar"
            if not session_path.exists():
                self.log("Matching .session file not found; scanning tokens in file...")
                try:
                    scanned_text, detected_format = FileHandler.read_file(file_path)
                except Exception as e:
                    self.log(f"Failed to scan file for tokens: {e}")
                    self.status_label.configure(text="File read error")
                    return

                token_candidates = self._extract_token_candidates(scanned_text)
                session_id, matched_count, ambiguous = (
                    self.vault.find_session_by_token_candidates(token_candidates))
                if ambiguous:
                    self.log("Multiple vault sessions match this file equally.")
                    self.log("Keep the original .session file beside the renamed file and try again.")
                    self.status_label.configure(text="Ambiguous session match")
                    return
                if not session_id:
                    self.log("No vault session matches the tokens in this file.")
                    self.log("Rehydration only works before the app closes or shreds the vault.")
                    self.status_label.configure(text="Session data not found")
                    return

                file_format = detected_format
                original_extension = original_path.suffix
                session_source = f"token scan ({matched_count} token match{'es' if matched_count != 1 else ''})"
                with open(session_path, 'w') as f:
                    json.dump({
                        'session_id': session_id,
                        'original_format': file_format,
                        'original_extension': original_extension,
                    }, f, indent=2)
            try:
                with open(session_path, 'r') as f:
                    content = f.read().strip()
                if content.startswith('{'):
                    metadata = json.loads(content)
                    session_id = metadata['session_id']
                    file_format = metadata.get('original_format', 'text')
                    original_extension = metadata.get('original_extension', Path(file_path).suffix)
                else:
                    session_id = content
                    file_format = 'text'
                    original_extension = Path(file_path).suffix
            except Exception as e:
                self.log(f"❌ Failed to read session file: {e}")
                self.status_label.configure(text="❌ Session file error")
                return

            mappings = self.vault.get_mapping(session_id)
            if not mappings:
                self.log("❌ No mappings found for this session.")
                self.status_label.configure(text="❌ Session data not found")
                return

            self.log(f"Session resolved by {session_source}: {session_id}")
            try:
                rehydrate_scan_text, _ = FileHandler.read_file(file_path)
            except Exception:
                rehydrate_scan_text = ""
            replacement_mappings = self._build_rehydration_map(mappings, rehydrate_scan_text)
            alias_count = len(replacement_mappings) - len(mappings)
            if alias_count:
                self.log(f"Recovered {alias_count} AI-modified token variant(s)")

            # --- Write output in original format ---
            output_path = original_path.parent / f"{original_path.stem}_RESTORED{original_extension}"
            restore_count = 0

            if file_format == 'pdf' and PYMUPDF_AVAILABLE:
                try:
                    restore_count = FileHandler.write_pdf_rehydrated(
                        file_path, str(output_path), replacement_mappings)
                    self.log(f"📄 PDF rehydrated in-place — {restore_count} token(s) restored")
                except Exception as e:
                    self.log(f"⚠️ PDF rehydration failed ({e}), falling back to .txt")
                    try:
                        sanitized_text, _ = FileHandler.read_file(file_path)
                    except Exception:
                        sanitized_text = ""
                    rehydrated_text = sanitized_text
                    for token, original_value in replacement_mappings.items():
                        if token in rehydrated_text:
                            rehydrated_text = rehydrated_text.replace(token, original_value)
                            restore_count += 1
                    output_path = original_path.parent / f"{original_path.stem}_RESTORED.txt"
                    FileHandler._write_text(output_path, rehydrated_text)
            elif file_format == 'pdf':
                # No PyMuPDF — extract text and do string replacement
                try:
                    sanitized_text, _ = FileHandler.read_file(file_path)
                except Exception as e:
                    self.log(f"❌ Failed to read sanitized file: {e}")
                    self.status_label.configure(text="❌ File read error")
                    return
                rehydrated_text = sanitized_text
                for token, original_value in replacement_mappings.items():
                    if token in rehydrated_text:
                        rehydrated_text = rehydrated_text.replace(token, original_value)
                        restore_count += 1
                output_path = original_path.parent / f"{original_path.stem}_RESTORED.txt"
                FileHandler._write_text(output_path, rehydrated_text)
                self.log("📄 Saved as .txt — run: pip install pymupdf  for PDF output")
            elif file_format == 'docx':
                try:
                    restore_count = FileHandler.write_docx_rehydrated(
                        file_path, str(output_path), replacement_mappings)
                    self.log(f"DOCX rehydrated in-place - {restore_count} token(s) restored")
                except Exception as e:
                    self.log(f"DOCX rehydration failed ({e}), falling back to .txt")
                    try:
                        sanitized_text, _ = FileHandler.read_file(file_path)
                    except Exception:
                        sanitized_text = ""
                    rehydrated_text = sanitized_text
                    for token, original_value in replacement_mappings.items():
                        if token in rehydrated_text:
                            rehydrated_text = rehydrated_text.replace(token, original_value)
                            restore_count += 1
                    output_path = original_path.parent / f"{original_path.stem}_RESTORED.txt"
                    FileHandler._write_text(output_path, rehydrated_text)
            else:
                # Non-PDF: plain text replacement
                try:
                    sanitized_text, _ = FileHandler.read_file(file_path)
                except Exception as e:
                    self.log(f"❌ Failed to read sanitized file: {e}")
                    self.status_label.configure(text="❌ File read error")
                    return
                rehydrated_text = sanitized_text
                for token, original_value in replacement_mappings.items():
                    if token in rehydrated_text:
                        rehydrated_text = rehydrated_text.replace(token, original_value)
                        restore_count += 1
                try:
                    FileHandler.write_file(str(output_path), rehydrated_text, file_format)
                except Exception as e:
                    self.log(f"⚠️ Could not write as {file_format.upper()}, saving as .txt")
                    output_path = original_path.parent / f"{original_path.stem}_RESTORED.txt"
                    FileHandler._write_text(output_path, rehydrated_text)

            # THE SHREDDER PROTOCOL
            self.vault.delete_session(session_id)
            deleted_session_files = self._delete_session_files_for_session(
                original_path.parent, session_id, session_path if session_path.exists() else None)

            self.log("✅ REHYDRATION COMPLETE")
            self.log(f"🔓 Restored {restore_count} sensitive items")
            self.log(f"💾 Restored file saved: {output_path.name}")
            self.log("🔥 THE SHREDDER PROTOCOL ACTIVATED")
            self.log(f"   • Session '{session_id}' deleted from vault")
            self.log(f"   • {deleted_session_files} session file(s) removed from disk")
            self.status_label.configure(
                text=f"✅ Restored {restore_count} items & shredded session"
            )

        except Exception as e:
            self.log(f"❌ Error: {e}")
            self.status_label.configure(text="❌ Rehydration failed")
            import traceback
            self.log(f"Debug: {traceback.format_exc()}")

    def on_closing(self):
        self.log("🔥 Application closing - initiating vault destruction...")
        self.vault.shred_vault()
        self.log("✅ Vault completely wiped. All session data destroyed.")
        self.destroy()


def main():
    app = PIISanitizerApp()
    app.mainloop()


if __name__ == "__main__":
    main()
